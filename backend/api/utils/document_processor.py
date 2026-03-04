"""
Document Processor Module
Xử lý các loại tài liệu: PDF, Word, Image (OCR)
"""
import os
from typing import Optional
from pypdf import PdfReader
from docx import Document
from PIL import Image
import pytesseract
from pytesseract import TesseractError, TesseractNotFoundError


tesseract_cmd = os.getenv("TESSERACT_CMD", "").strip()
if tesseract_cmd:
    pytesseract.pytesseract.tesseract_cmd = tesseract_cmd


def extract_text_from_pdf(file_path: str) -> str:
    """
    Trích xuất text từ file PDF sử dụng PyPDF
    """
    try:
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n\n".join(text_parts)
    except Exception as e:
        raise Exception(f"Lỗi khi đọc PDF: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """
    Trích xuất text từ file Word (.docx) sử dụng python-docx
    """
    try:
        doc = Document(file_path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Cũng trích xuất text từ tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts)
    except Exception as e:
        raise Exception(f"Lỗi khi đọc Word: {str(e)}")


def extract_text_from_image(file_path: str, lang: str = 'vie+eng') -> str:
    """
    Trích xuất text từ ảnh sử dụng Tesseract OCR
    Hỗ trợ tiếng Việt và tiếng Anh
    
    Args:
        file_path: Đường dẫn đến file ảnh
        lang: Ngôn ngữ OCR (mặc định: vie+eng cho tiếng Việt + Anh)
    """
    try:
        available_langs = set(pytesseract.get_languages(config=''))
        requested_langs = [code for code in lang.split('+') if code]

        valid_langs = [code for code in requested_langs if code in available_langs]
        if not valid_langs:
            if 'eng' in available_langs:
                valid_langs = ['eng']
            elif available_langs:
                valid_langs = [sorted(available_langs)[0]]
            else:
                raise Exception(
                    "Không tìm thấy language data cho Tesseract. "
                    "Hãy cài gói ngôn ngữ (ví dụ: tesseract-ocr-vie, tesseract-ocr-eng)."
                )

        image = Image.open(file_path)
        # Convert to RGB if necessary
        if image.mode != 'RGB':
            image = image.convert('RGB')

        text = pytesseract.image_to_string(
            image,
            lang='+'.join(valid_langs),
            config='--oem 3 --psm 6'
        )
        return text.strip()
    except TesseractNotFoundError:
        raise Exception(
            "Chưa cài Tesseract OCR trên máy chủ hoặc chưa thêm vào PATH. "
            "Vui lòng cài `tesseract-ocr` và các gói ngôn ngữ cần thiết."
        )
    except TesseractError as e:
        raise Exception(f"Lỗi Tesseract OCR: {str(e)}")
    except Exception as e:
        raise Exception(f"Lỗi khi OCR ảnh: {str(e)}")


def process_document(file_path: str, file_type: str) -> str:
    """
    Xử lý tài liệu dựa trên loại file
    
    Args:
        file_path: Đường dẫn đến file
        file_type: Loại file (pdf, docx, png, jpg, jpeg)
    
    Returns:
        Text đã trích xuất từ tài liệu
    """
    file_type = file_type.lower()
    
    if file_type == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_type in ['docx', 'doc']:
        return extract_text_from_docx(file_path)
    elif file_type in ['png', 'jpg', 'jpeg', 'bmp', 'tiff']:
        return extract_text_from_image(file_path)
    else:
        raise ValueError(f"Không hỗ trợ định dạng file: {file_type}")


def get_file_extension(filename: str) -> str:
    """Lấy extension của file"""
    return filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
